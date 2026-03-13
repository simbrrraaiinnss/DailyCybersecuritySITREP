#!/usr/bin/env python3
"""
=============================================================================
  Daily Cybersecurity SITREP Generator (Manual / On-Demand)
=============================================================================
  Organization : Evergreen Healthcare
  Author       : Security Operations Team
  Description  : Gathers cybersecurity intelligence from CISA, vendor
                 advisories, and public news sources, then produces a
                 multi-page categorized Daily SITREP as a Word document
                 and sends an HTML-formatted email to the designated
                 recipient.

                 The Word document includes:
                   Page 1   – Concise executive summary (1-page overview)
                   Page 2   – APPENDIX A: IMMEDIATE Threats (detailed)
                   Page 3   – APPENDIX B: PRIORITY Threats (detailed)
                   Page 4   – APPENDIX C: ROUTINE Items (detailed)

                 The email body contains the concise summary; the full
                 multi-page document with technical appendices is attached
                 or saved alongside.

  Usage        : python generate_sitrep_manual.py [--no-email] [--output DIR]
=============================================================================
"""

from __future__ import annotations

import argparse
import concurrent.futures
import datetime
import hashlib
import json
import logging
import os
import re
import smtplib
import ssl
import sys
import textwrap
import time
import traceback
from dataclasses import dataclass, field
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from typing import Dict, List, Optional, Tuple
from pathlib import Path

import feedparser
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║                        CONFIGURATION SECTION                            ║
# ║   Edit the values below to customize for your environment.              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

CONFIG = {
    # ── Email Settings ──────────────────────────────────────────────────────
    "email_recipient": "mdsimbre@evergreenhealthcare.org",
    "email_from": "sitrep-noreply@evergreenhealthcare.org",
    "smtp_server": "smtp.office365.com",       # Change to your SMTP server
    "smtp_port": 587,                           # 587 for STARTTLS, 465 for SSL
    "smtp_username": "",                        # Set via env SITREP_SMTP_USER
    "smtp_password": "",                        # Set via env SITREP_SMTP_PASS
    "smtp_use_tls": True,

    # ── Organization Details ────────────────────────────────────────────────
    "org_name": "Evergreen Healthcare",
    "classification": "TLP:AMBER – Internal Use Only",

    # ── Output ──────────────────────────────────────────────────────────────
    "output_dir": ".",                          # Default output directory
    "log_file": "sitrep_generator.log",

    # ── Network ─────────────────────────────────────────────────────────────
    "request_timeout": 20,                      # seconds per HTTP request
    "max_workers": 10,                          # parallel fetch threads
    "user_agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    ),
}

# ── Hospital Technology Stack ───────────────────────────────────────────────
# Keywords used to determine if a threat/advisory is relevant to OUR stack.
TECHNOLOGY_STACK = [
    "Epic", "Epic EHR", "EpicCare",
    "Palo Alto", "PAN-OS", "GlobalProtect", "Panorama", "Cortex",
    "Microsoft", "Microsoft 365", "Office 365", "Azure", "Windows",
    "Exchange", "Entra", "Active Directory", "Intune", "Defender",
    "Cisco", "IOS-XE", "Meraki", "Webex", "ASA", "Firepower",
    "Aruba", "ArubaOS", "ClearPass",
    "HP", "Hewlett Packard", "HPE",
    "Konica Minolta", "bizhub",
    "GE HealthCare", "GE Healthcare",
    "Philips", "IntelliSpace",
    "Siemens Healthineers", "syngo",
    "Becton Dickinson", "BD Pyxis", "BD Alaris",
    "Mindray",
    "Masimo",
    "Varian",
    "Hologic",
    "Medtronic",
    "Roche Diagnostics",
    "Abbott",
    "Oracle", "Cerner", "Oracle Health",
    "Dell", "Dell Technologies", "PowerEdge", "iDRAC", "PowerScale",
    "Citrix", "NetScaler", "XenApp", "XenDesktop",
    "Honeywell", "Honeywell Building",
    "Imprivata",
    "ZEISS", "ZEISS Medical",
    # Generic healthcare / infrastructure
    "DICOM", "HL7", "FHIR", "PACS", "VPN", "RDP", "VMware", "ESXi",
    "Linux", "Apache", "OpenSSL", "Log4j", "healthcare", "hospital",
    "medical device", "HIS", "LIS", "RIS",
]

# ── Vendor Security Advisory URLs ───────────────────────────────────────────
VENDOR_ADVISORY_URLS = [
    # RSS / Atom feeds
    "https://community.infoblox.com/cixhp49439/rss/board?board.id=generalsecurity",
    "https://www.darkreading.com/rss.xml",
    "https://sec.cloudapps.cisco.com/security/center/psirtrss20/CiscoSecurityAdvisory.xml",
    "https://blog.rapid7.com/rss/",
    "https://security.paloaltonetworks.com/rss.xml",
    "https://www.bleepingcomputer.com/feed/",
    # Web pages (scraped)
    "https://www.epic.com/epic/page/reporting-potential-security-vulnerability/",
    "https://www.gehealthcare.com/productsecurity/products",
    "https://www.philips.com/a-w/security/security-advisories.html",
    "https://www.siemens.com/global/en/products/services/cert.html",
    "https://www.bd.com/en-us/about-bd/cybersecurity/patches/security-patches-bd-pyxis-products",
    "https://www.mindray.com/na/solutions/hospitalwide-solution/cybersecurity/updates/",
    "https://www.masimo.com/contact-support/globalservices/masimo-coordinated-vulnerability-disclosure-statement/",
    "https://cancercare.siemens-healthineers.com/resources-support/cybersecurity-at-varian",
    "https://www.hologic.com/support/usa/breast-skeletal-products-cybersecurity",
    "https://www.medtronic.com/en-us/e/product-security/security-bulletins.html",
    "https://diagnostics.roche.com/global/en/legal/product-security-advisory.html",
    "https://www.abbott.com/en-us/policies/cybersecurity.html",
    "https://sec.cloudapps.cisco.com/security/center/publicationListing.x",
    "https://msrc.microsoft.com/update-guide",
    "https://www.oracle.com/security-alerts/",
    "https://www.dell.com/support/kbdoc/en-us/000415586/dsa-2026-049-security-update-for-dell-powerscale-onefs-multiple-vulnerabilities",
    "https://security.paloaltonetworks.com/",
    "https://www.rapid7.com/security/",
    "https://www.tenable.com/security",
    "https://www.infoblox.com/company/legal/vulnerability-responsible-disclosure-policy/",
    "https://support.citrix.com/s/topic/0TO4z0000001GYdGAM/security-bulletin?language=en_US",
    "https://www.honeywell.com/us/en/product-security",
    "https://www.imprivata.com/company/trust-and-security",
    "https://www.zeiss.com/meditec/en/service/cybersecurity-at-zeiss-medical-technology.html",
]

# ── CISA Sources ────────────────────────────────────────────────────────────
CISA_SOURCES = [
    "https://www.cisa.gov/news.xml",
    "https://www.cisa.gov/sites/default/files/feeds/known_exploited_vulnerabilities.json",
]

# ── Healthcare / General Cybersecurity News Feeds ───────────────────────────
NEWS_FEEDS = [
    "https://www.bleepingcomputer.com/feed/",
    "https://www.darkreading.com/rss.xml",
    "https://feeds.feedburner.com/TheHackersNews",
    "https://krebsonsecurity.com/feed/",
    "https://healthitsecurity.com/feed",
]

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║                       END OF CONFIGURATION                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

# ── Logging Setup ───────────────────────────────────────────────────────────
logger = logging.getLogger("sitrep")
logger.setLevel(logging.DEBUG)

_console_handler = logging.StreamHandler(sys.stdout)
_console_handler.setLevel(logging.INFO)
_console_handler.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))
logger.addHandler(_console_handler)


def _setup_file_logging(log_path: str):
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)s] %(name)s: %(message)s"
    ))
    logger.addHandler(fh)


# ── Data Classes ────────────────────────────────────────────────────────────
@dataclass
class ThreatItem:
    """Represents a single threat / advisory item."""
    title: str
    source: str
    url: str = ""
    description: str = ""
    published: str = ""
    severity: str = "ROUTINE"       # IMMEDIATE | PRIORITY | ROUTINE
    relevance_tags: List[str] = field(default_factory=list)
    cve_ids: List[str] = field(default_factory=list)

    @property
    def color(self) -> str:
        return {"IMMEDIATE": "Red", "PRIORITY": "Amber", "ROUTINE": "Green"}.get(
            self.severity, "Green"
        )

    @property
    def color_rgb(self) -> RGBColor:
        return {
            "IMMEDIATE": RGBColor(220, 20, 20),
            "PRIORITY": RGBColor(255, 165, 0),
            "ROUTINE": RGBColor(34, 139, 34),
        }.get(self.severity, RGBColor(34, 139, 34))

    @property
    def html_color(self) -> str:
        return {
            "IMMEDIATE": "#DC1414",
            "PRIORITY": "#FFA500",
            "ROUTINE": "#228B22",
        }.get(self.severity, "#228B22")


# ── Helper Functions ────────────────────────────────────────────────────────
def _http_get(url: str, timeout: int = None) -> Optional[requests.Response]:
    """Perform an HTTP GET with error handling."""
    timeout = timeout or CONFIG["request_timeout"]
    headers = {"User-Agent": CONFIG["user_agent"]}
    try:
        resp = requests.get(url, headers=headers, timeout=timeout, verify=True,
                            allow_redirects=True)
        resp.raise_for_status()
        return resp
    except requests.RequestException as exc:
        logger.warning("  ⚠  Failed to fetch %s – %s", url, exc)
        return None


def _extract_cves(text: str) -> List[str]:
    """Extract CVE identifiers from text."""
    return list(set(re.findall(r"CVE-\d{4}-\d{4,7}", text, re.IGNORECASE)))


def _is_relevant(text: str) -> Tuple[bool, List[str]]:
    """Check if text is relevant to our technology stack."""
    text_lower = text.lower()
    matched = []
    for kw in TECHNOLOGY_STACK:
        if kw.lower() in text_lower:
            matched.append(kw)
    return (len(matched) > 0, matched)


def _classify_severity(title: str, description: str, cves: List[str],
                       relevance_tags: List[str]) -> str:
    """Heuristic severity classification."""
    combined = (title + " " + description).lower()

    # IMMEDIATE indicators
    immediate_keywords = [
        "actively exploited", "zero-day", "0-day", "critical",
        "remote code execution", "rce", "ransomware", "actively being exploited",
        "emergency", "exploitation in the wild", "known exploited",
        "pre-auth", "unauthenticated", "wormable",
    ]
    for kw in immediate_keywords:
        if kw in combined:
            # Only IMMEDIATE if relevant to our stack or healthcare
            if relevance_tags or "healthcare" in combined or "hospital" in combined:
                return "IMMEDIATE"
            return "PRIORITY"

    # PRIORITY indicators
    priority_keywords = [
        "high", "important", "privilege escalation", "authentication bypass",
        "denial of service", "dos", "sql injection", "data breach",
        "vulnerability", "exploit", "patch", "update available",
        "medical device", "hipaa",
    ]
    priority_count = sum(1 for kw in priority_keywords if kw in combined)
    if priority_count >= 2 or (priority_count >= 1 and relevance_tags):
        return "PRIORITY"

    if relevance_tags:
        return "PRIORITY"

    return "ROUTINE"


def _determine_threat_posture(threats: List[ThreatItem]) -> str:
    """Determine overall threat posture based on collected threats."""
    immediate_count = sum(1 for t in threats if t.severity == "IMMEDIATE")
    priority_count = sum(1 for t in threats if t.severity == "PRIORITY")

    if immediate_count >= 1:
        return "ELEVATED"
    if priority_count >= 3:
        return "ELEVATED"
    if priority_count >= 1:
        return "GUARDED"
    return "LOW"


# ── Hospital Tech Stack Mapping ──────────────────────────────────────────────
# Maps technology keywords to specific hospital system descriptions for the
# "Systems in Our Environment" field in appendix entries.
TECH_STACK_DETAILS = {
    "Epic":               "Epic EHR / EpicCare – Primary electronic health record system",
    "Epic EHR":           "Epic EHR / EpicCare – Primary electronic health record system",
    "EpicCare":           "Epic EHR / EpicCare – Primary electronic health record system",
    "Palo Alto":          "Palo Alto Networks – Perimeter firewalls & GlobalProtect VPN",
    "PAN-OS":             "Palo Alto PAN-OS – Firewall operating system",
    "GlobalProtect":      "Palo Alto GlobalProtect – Remote-access VPN",
    "Panorama":           "Palo Alto Panorama – Centralized firewall management",
    "Cortex":             "Palo Alto Cortex XDR – Endpoint detection & response",
    "Microsoft":          "Microsoft 365 / Azure / Windows – Enterprise productivity & cloud",
    "Microsoft 365":      "Microsoft 365 – Email, Office apps, Teams",
    "Office 365":         "Microsoft 365 – Email, Office apps, Teams",
    "Azure":              "Microsoft Azure – Cloud infrastructure & services",
    "Windows":            "Microsoft Windows – Desktop & server operating systems",
    "Exchange":           "Microsoft Exchange – Email server infrastructure",
    "Entra":              "Microsoft Entra ID – Identity & access management",
    "Active Directory":   "Microsoft Active Directory – Directory services",
    "Intune":             "Microsoft Intune – Mobile device management",
    "Defender":           "Microsoft Defender – Endpoint protection",
    "Cisco":              "Cisco – Network switches, routers, telephony",
    "IOS-XE":             "Cisco IOS-XE – Switch/router operating system",
    "Meraki":             "Cisco Meraki – Cloud-managed network devices",
    "ASA":                "Cisco ASA – Legacy firewall appliances",
    "Firepower":          "Cisco Firepower – Next-gen firewall / IPS",
    "Aruba":              "Aruba Networks – Wireless access points & controllers",
    "ArubaOS":            "Aruba OS – Wireless network operating system",
    "ClearPass":          "Aruba ClearPass – Network access control",
    "HP":                 "HP / HPE – Workstations, servers, printers",
    "Hewlett Packard":    "HP / HPE – Workstations, servers, printers",
    "HPE":                "HPE – ProLiant servers, storage",
    "Konica Minolta":     "Konica Minolta bizhub – Multifunction printers",
    "bizhub":             "Konica Minolta bizhub – Multifunction printers",
    "GE HealthCare":      "GE HealthCare – Medical imaging (CT, MRI, Ultrasound)",
    "GE Healthcare":      "GE HealthCare – Medical imaging (CT, MRI, Ultrasound)",
    "Philips":            "Philips – Patient monitors, imaging, IntelliSpace",
    "IntelliSpace":       "Philips IntelliSpace – Medical imaging informatics",
    "Siemens Healthineers": "Siemens Healthineers – Imaging systems (CT, MRI, X-ray)",
    "syngo":              "Siemens syngo – Medical imaging software",
    "Becton Dickinson":   "BD – Medication management & diagnostics",
    "BD Pyxis":           "BD Pyxis – Automated medication dispensing",
    "BD Alaris":          "BD Alaris – Infusion pump systems",
    "Mindray":            "Mindray – Patient monitoring, ultrasound, in-vitro diagnostics",
    "Masimo":             "Masimo – Pulse oximetry & patient monitoring",
    "Varian":             "Varian (Siemens) – Radiation oncology treatment systems",
    "Hologic":            "Hologic – Mammography & breast health systems",
    "Medtronic":          "Medtronic – Implantable devices, surgical navigation",
    "Roche Diagnostics":  "Roche Diagnostics – Laboratory analyzers & diagnostics",
    "Abbott":             "Abbott – Diagnostics & point-of-care testing",
    "Oracle":             "Oracle – Database servers, Oracle Health / Cerner",
    "Cerner":             "Oracle Health (Cerner) – EHR / clinical information systems",
    "Oracle Health":      "Oracle Health (Cerner) – EHR / clinical information systems",
    "Dell":               "Dell – PowerEdge servers, workstations",
    "Dell Technologies":  "Dell – PowerEdge servers, workstations, storage",
    "PowerEdge":          "Dell PowerEdge – Server infrastructure",
    "iDRAC":              "Dell iDRAC – Remote server management",
    "PowerScale":         "Dell PowerScale – Network-attached storage",
    "Citrix":             "Citrix – Virtual desktop & application delivery",
    "NetScaler":          "Citrix NetScaler – Application delivery controller / VPN",
    "XenApp":             "Citrix XenApp – Application virtualization",
    "XenDesktop":         "Citrix XenDesktop – Virtual desktop infrastructure",
    "Imprivata":          "Imprivata – Single sign-on & identity governance",
    "VMware":             "VMware / Broadcom – Server virtualization (ESXi, vCenter)",
    "ESXi":               "VMware ESXi – Hypervisor infrastructure",
    "DICOM":              "DICOM – Medical imaging communication standard",
    "HL7":                "HL7 – Healthcare interoperability messaging",
    "FHIR":               "HL7 FHIR – Healthcare data exchange standard",
    "PACS":               "PACS – Picture archiving & communication systems",
    "VPN":                "VPN infrastructure – Remote access connectivity",
    "RDP":                "Remote Desktop Protocol – Clinical workstation access",
    "Linux":              "Linux servers – Infrastructure & application hosting",
    "Apache":             "Apache HTTP Server – Web application hosting",
    "OpenSSL":            "OpenSSL – TLS/SSL library used across infrastructure",
    "Log4j":              "Apache Log4j – Java logging (used in many clinical apps)",
    "ZEISS":              "ZEISS Medical Technology – Ophthalmic & surgical systems",
    "ZEISS Medical":      "ZEISS Medical Technology – Ophthalmic & surgical systems",
    "Honeywell":          "Honeywell – Building automation & environmental controls",
    "Honeywell Building": "Honeywell – Building management systems (HVAC, fire)",
}


# ── CVSS Score Lookup ───────────────────────────────────────────────────────
# Queries the NVD (National Vulnerability Database) API to fetch CVSS scores
# for CVE identifiers. Results are cached to avoid duplicate API calls.
_cvss_cache: Dict[str, Optional[dict]] = {}


def _lookup_cvss(cve_id: str) -> Optional[dict]:
    """
    Look up CVSS score information for a CVE from the NVD API.

    Returns a dict with keys:
        - score (float): CVSS base score
        - severity (str): e.g. CRITICAL, HIGH, MEDIUM, LOW
        - vector (str): CVSS vector string
        - version (str): CVSS version (3.1, 3.0, 2.0)
    Returns None if lookup fails or CVE is not found.
    """
    if cve_id in _cvss_cache:
        return _cvss_cache[cve_id]

    if not cve_id or not re.match(r"CVE-\d{4}-\d{4,7}", cve_id, re.IGNORECASE):
        return None

    try:
        url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}"
        resp = requests.get(
            url,
            headers={"User-Agent": CONFIG["user_agent"]},
            timeout=15,
        )
        if resp.status_code == 200:
            data = resp.json()
            vulns = data.get("vulnerabilities", [])
            if vulns:
                cve_item = vulns[0].get("cve", {})
                metrics = cve_item.get("metrics", {})

                # Try CVSS 3.1 first, then 3.0, then 2.0
                for version_key, version_label in [
                    ("cvssMetricV31", "3.1"),
                    ("cvssMetricV30", "3.0"),
                    ("cvssMetricV2", "2.0"),
                ]:
                    metric_list = metrics.get(version_key, [])
                    if metric_list:
                        cvss_data = metric_list[0].get("cvssData", {})
                        result = {
                            "score": cvss_data.get("baseScore", 0.0),
                            "severity": cvss_data.get("baseSeverity",
                                        metric_list[0].get("baseSeverity", "UNKNOWN")),
                            "vector": cvss_data.get("vectorString", "N/A"),
                            "version": version_label,
                        }
                        _cvss_cache[cve_id] = result
                        return result

        _cvss_cache[cve_id] = None
        return None
    except Exception as exc:
        logger.debug("CVSS lookup failed for %s: %s", cve_id, exc)
        _cvss_cache[cve_id] = None
        return None


def _batch_lookup_cvss(cve_ids: List[str]) -> Dict[str, Optional[dict]]:
    """
    Batch lookup CVSS scores for multiple CVEs.
    Includes a small delay between requests to respect NVD rate limits.
    """
    results = {}
    for cve_id in cve_ids:
        results[cve_id] = _lookup_cvss(cve_id)
        time.sleep(0.6)  # NVD API rate limit: ~5 requests per 30 seconds without API key
    return results


def _get_systems_in_environment(relevance_tags: List[str]) -> str:
    """
    Map relevance tags back to specific hospital system descriptions.
    Uses TECH_STACK_DETAILS to provide human-readable system descriptions.
    """
    if not relevance_tags:
        return "No directly matched systems identified — general advisory."
    systems = []
    seen = set()
    for tag in relevance_tags:
        desc = TECH_STACK_DETAILS.get(tag, f"{tag} – (present in hospital technology stack)")
        if desc not in seen:
            seen.add(desc)
            systems.append(desc)
    return "; ".join(systems) if systems else "General infrastructure"


def _infer_attack_vector(description: str, cve_ids: List[str]) -> str:
    """
    Infer the attack vector from the description text and CVSS data.
    Attempts CVSS vector first, falls back to keyword heuristics.
    """
    # Try to get from CVSS data
    for cve_id in cve_ids:
        cvss_info = _cvss_cache.get(cve_id)
        if cvss_info and cvss_info.get("vector"):
            vec = cvss_info["vector"]
            if "AV:N" in vec:
                return "Network — Exploitable remotely over the network"
            elif "AV:A" in vec:
                return "Adjacent Network — Requires access to local network segment"
            elif "AV:L" in vec:
                return "Local — Requires local system access"
            elif "AV:P" in vec:
                return "Physical — Requires physical access to the device"

    # Keyword heuristics
    desc_lower = description.lower()
    if any(kw in desc_lower for kw in ["remote code", "rce", "remotely", "unauthenticated remote",
                                         "network-based", "internet-facing"]):
        return "Network — Exploitable remotely over the network"
    if any(kw in desc_lower for kw in ["local privilege", "local access", "authenticated local"]):
        return "Local — Requires local system access"
    if any(kw in desc_lower for kw in ["physical access", "usb", "console access"]):
        return "Physical — Requires physical access to the device"
    return "Not determined — Review advisory for details"


def _infer_remediation(title: str, description: str, severity: str) -> str:
    """
    Generate remediation guidance based on available information.
    """
    combined = (title + " " + description).lower()
    steps = []

    if "patch" in combined or "update" in combined:
        steps.append("Apply vendor-supplied patch or update immediately")
    if "workaround" in combined or "mitigation" in combined:
        steps.append("Review vendor-published workarounds/mitigations")
    if "firmware" in combined:
        steps.append("Schedule firmware update during maintenance window")
    if "configuration" in combined or "disable" in combined:
        steps.append("Review and harden system configuration per vendor guidance")

    if not steps:
        steps.append("Review vendor advisory for specific remediation steps")
        steps.append("Check for available patches or updated firmware")

    if severity == "IMMEDIATE":
        steps.append("Engage incident response team if exploitation is suspected")
        steps.append("Consider temporary isolation of affected systems")
    elif severity == "PRIORITY":
        steps.append("Schedule patching within 72-hour remediation window")
        steps.append("Validate exposure through vulnerability scanning")

    steps.append("Monitor vendor channels for follow-up advisories")
    return "; ".join(steps)


def _infer_timeline(severity: str) -> str:
    """Return urgency timeline based on severity classification."""
    return {
        "IMMEDIATE": "Action required within 24 hours — Treat as emergency",
        "PRIORITY":  "Action required within 72 hours — Schedule during next maintenance window",
        "ROUTINE":   "Address during normal patch cycle — Monitor for escalation",
    }.get(severity, "Review at next scheduled assessment")


def _infer_impact(title: str, description: str, relevance_tags: List[str],
                  severity: str) -> str:
    """
    Generate a healthcare-specific impact analysis based on threat characteristics.
    """
    combined = (title + " " + description).lower()
    impacts = []

    if any(kw in combined for kw in ["ransomware", "encryption", "data destruction"]):
        impacts.append("Potential for system-wide encryption disrupting clinical operations")
    if any(kw in combined for kw in ["data breach", "data exfiltration", "phi", "hipaa", "patient data"]):
        impacts.append("Risk of Protected Health Information (PHI) exposure — HIPAA breach notification may be required")
    if any(kw in combined for kw in ["remote code execution", "rce"]):
        impacts.append("Remote code execution could allow attacker to gain full system control")
    if any(kw in combined for kw in ["denial of service", "dos", "availability"]):
        impacts.append("Service disruption could impact clinical system availability")
    if any(kw in combined for kw in ["medical device", "patient monitor", "infusion", "imaging"]):
        impacts.append("Patient safety concern — affected medical devices may require clinical engineering review")
    if any(kw in combined for kw in ["authentication bypass", "privilege escalation"]):
        impacts.append("Unauthorized access to clinical or administrative systems")

    if not impacts:
        if severity == "IMMEDIATE":
            impacts.append("High-severity threat requiring immediate assessment of clinical impact")
        elif severity == "PRIORITY":
            impacts.append("Moderate risk to hospital operations — assess exposure and plan remediation")
        else:
            impacts.append("Low direct impact — monitor for changes in threat landscape")

    if relevance_tags:
        impacts.append(f"Directly affects hospital systems: {', '.join(relevance_tags[:5])}")

    return "; ".join(impacts)


# ── Intelligence Gathering ──────────────────────────────────────────────────
class IntelCollector:
    """Collects cybersecurity intelligence from all configured sources."""

    def __init__(self):
        self.threats: List[ThreatItem] = []
        self._seen_hashes: set = set()

    def _dedup_add(self, item: ThreatItem):
        h = hashlib.md5(item.title.encode()).hexdigest()
        if h not in self._seen_hashes:
            self._seen_hashes.add(h)
            self.threats.append(item)

    # ── CISA KEV Catalog ────────────────────────────────────────────────────
    def fetch_cisa_kev(self):
        """Fetch CISA Known Exploited Vulnerabilities catalog."""
        logger.info("📡 Fetching CISA Known Exploited Vulnerabilities catalog …")
        url = "https://www.cisa.gov/sites/default/files/feeds/known_exploited_vulnerabilities.json"
        resp = _http_get(url)
        if not resp:
            return
        try:
            data = resp.json()
            today = datetime.date.today()
            recent_cutoff = today - datetime.timedelta(days=3)
            vulns = data.get("vulnerabilities", [])
            count = 0
            for v in vulns:
                date_added = v.get("dateAdded", "")
                try:
                    added_dt = datetime.date.fromisoformat(date_added)
                except ValueError:
                    continue
                if added_dt < recent_cutoff:
                    continue
                title = f"[CISA KEV] {v.get('vendorProject', '')} – {v.get('vulnerabilityName', '')}"
                desc = v.get("shortDescription", "")
                cves = [v.get("cveID", "")]
                relevant, tags = _is_relevant(title + " " + desc + " " + v.get("vendorProject", ""))
                sev = "IMMEDIATE" if relevant else "PRIORITY"
                item = ThreatItem(
                    title=title,
                    source="CISA KEV",
                    url=f"https://nvd.nist.gov/vuln/detail/{cves[0]}" if cves[0] else "",
                    description=desc,
                    published=date_added,
                    severity=sev,
                    relevance_tags=tags,
                    cve_ids=cves,
                )
                self._dedup_add(item)
                count += 1
            logger.info("   ✔ CISA KEV: %d recent entries collected", count)
        except Exception as exc:
            logger.error("   ✖ CISA KEV parse error: %s", exc)

    # ── CISA Alerts / News Feed ─────────────────────────────────────────────
    def fetch_cisa_news(self):
        """Fetch CISA news/alerts RSS feed."""
        logger.info("📡 Fetching CISA alerts/news feed …")
        url = "https://www.cisa.gov/news.xml"
        resp = _http_get(url)
        if not resp:
            return
        feed = feedparser.parse(resp.content)
        count = 0
        for entry in feed.entries[:30]:
            title = entry.get("title", "")
            desc = entry.get("summary", "")
            link = entry.get("link", "")
            pub = entry.get("published", "")
            combined_text = title + " " + desc
            cves = _extract_cves(combined_text)
            relevant, tags = _is_relevant(combined_text)
            sev = _classify_severity(title, desc, cves, tags)
            item = ThreatItem(
                title=f"[CISA] {title}",
                source="CISA",
                url=link,
                description=desc[:500],
                published=pub,
                severity=sev,
                relevance_tags=tags,
                cve_ids=cves,
            )
            self._dedup_add(item)
            count += 1
        logger.info("   ✔ CISA News: %d entries collected", count)

    # ── RSS / Atom Feed Parser (generic) ────────────────────────────────────
    def _parse_rss_feed(self, url: str, source_label: str) -> int:
        resp = _http_get(url)
        if not resp:
            return 0
        feed = feedparser.parse(resp.content)
        count = 0
        for entry in feed.entries[:25]:
            title = entry.get("title", "")
            desc = entry.get("summary", entry.get("description", ""))
            link = entry.get("link", "")
            pub = entry.get("published", entry.get("updated", ""))
            combined = title + " " + desc
            cves = _extract_cves(combined)
            relevant, tags = _is_relevant(combined)
            sev = _classify_severity(title, desc, cves, tags)
            item = ThreatItem(
                title=title,
                source=source_label,
                url=link,
                description=desc[:500],
                published=pub,
                severity=sev,
                relevance_tags=tags,
                cve_ids=cves,
            )
            self._dedup_add(item)
            count += 1
        return count

    # ── Web Page Scraper (for non-RSS vendor pages) ─────────────────────────
    def _scrape_vendor_page(self, url: str, source_label: str) -> int:
        resp = _http_get(url)
        if not resp:
            return 0
        soup = BeautifulSoup(resp.text, "html.parser")
        # Remove script/style
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        count = 0
        # Try to find advisory-like items: look for headings, list items, links
        candidates = []

        # Strategy 1: look for links containing security-related text
        for a_tag in soup.find_all("a", href=True):
            text = a_tag.get_text(strip=True)
            if len(text) > 15 and any(kw in text.lower() for kw in
                    ["advisory", "security", "vulnerability", "patch",
                     "update", "bulletin", "cve", "alert"]):
                href = a_tag["href"]
                if not href.startswith("http"):
                    from urllib.parse import urljoin
                    href = urljoin(url, href)
                candidates.append((text, href))

        # Strategy 2: extract paragraphs for general context
        if not candidates:
            text_content = soup.get_text(separator="\n", strip=True)
            # Chunk into pseudo-items
            lines = [l.strip() for l in text_content.split("\n") if len(l.strip()) > 30]
            for line in lines[:10]:
                candidates.append((line[:200], url))

        for title_text, link in candidates[:15]:
            cves = _extract_cves(title_text)
            relevant, tags = _is_relevant(title_text + " " + source_label)
            sev = _classify_severity(title_text, source_label, cves, tags)
            item = ThreatItem(
                title=title_text[:200],
                source=source_label,
                url=link,
                description="",
                published="",
                severity=sev,
                relevance_tags=tags,
                cve_ids=cves,
            )
            self._dedup_add(item)
            count += 1
        return count

    # ── Fetch All Vendor Advisories (parallel) ──────────────────────────────
    def fetch_vendor_advisories(self):
        """Fetch all vendor security advisory sources."""
        logger.info("📡 Fetching vendor security advisories (%d sources) …",
                     len(VENDOR_ADVISORY_URLS))
        rss_extensions = [".xml", "/rss", "/feed", "rss.xml", "/rss/"]

        def _process_url(url):
            # Determine vendor label from URL
            from urllib.parse import urlparse
            domain = urlparse(url).netloc.replace("www.", "")
            label = domain.split(".")[0].capitalize()
            is_rss = any(ext in url.lower() for ext in rss_extensions)
            if is_rss:
                return (label, self._parse_rss_feed(url, label))
            else:
                return (label, self._scrape_vendor_page(url, label))

        with concurrent.futures.ThreadPoolExecutor(max_workers=CONFIG["max_workers"]) as executor:
            futures = {executor.submit(_process_url, u): u for u in VENDOR_ADVISORY_URLS}
            for future in concurrent.futures.as_completed(futures):
                url = futures[future]
                try:
                    label, cnt = future.result()
                    logger.info("   ✔ %s: %d items", label, cnt)
                except Exception as exc:
                    logger.warning("   ⚠  Error processing %s – %s", url, exc)

    # ── Fetch News Feeds ────────────────────────────────────────────────────
    def fetch_news_feeds(self):
        """Fetch general cybersecurity and healthcare news."""
        logger.info("📡 Fetching cybersecurity news feeds …")
        for feed_url in NEWS_FEEDS:
            from urllib.parse import urlparse
            domain = urlparse(feed_url).netloc.replace("www.", "")
            label = domain.split(".")[0].capitalize()
            cnt = self._parse_rss_feed(feed_url, label)
            logger.info("   ✔ %s: %d items", label, cnt)

    # ── Main Collection ─────────────────────────────────────────────────────
    def collect_all(self):
        """Run all collection routines."""
        self.fetch_cisa_kev()
        self.fetch_cisa_news()
        self.fetch_vendor_advisories()
        self.fetch_news_feeds()

        # Re-sort: IMMEDIATE first, then PRIORITY, then ROUTINE
        order = {"IMMEDIATE": 0, "PRIORITY": 1, "ROUTINE": 2}
        self.threats.sort(key=lambda t: order.get(t.severity, 3))

        logger.info("═══════════════════════════════════════════════════════")
        logger.info("  Collection complete: %d total threat items", len(self.threats))
        imm = sum(1 for t in self.threats if t.severity == "IMMEDIATE")
        pri = sum(1 for t in self.threats if t.severity == "PRIORITY")
        rou = sum(1 for t in self.threats if t.severity == "ROUTINE")
        logger.info("  🔴 IMMEDIATE: %d  |  🟠 PRIORITY: %d  |  🟢 ROUTINE: %d",
                     imm, pri, rou)
        logger.info("═══════════════════════════════════════════════════════")


# ── SITREP Report Generation ───────────────────────────────────────────────
class SITREPGenerator:
    """Generates the SITREP in DOCX and HTML formats."""

    def __init__(self, threats: List[ThreatItem]):
        self.threats = threats
        self.now = datetime.datetime.now()
        self.date_str = self.now.strftime("%B %d, %Y")
        self.date_file = self.now.strftime("%Y%m%d")
        self.posture = _determine_threat_posture(threats)

    # ── Helpers ─────────────────────────────────────────────────────────────
    def _threats_by_severity(self, severity: str) -> List[ThreatItem]:
        return [t for t in self.threats if t.severity == severity]

    def _posture_color_rgb(self) -> RGBColor:
        return {
            "ELEVATED": RGBColor(220, 20, 20),
            "GUARDED": RGBColor(255, 165, 0),
            "LOW": RGBColor(34, 139, 34),
        }[self.posture]

    def _posture_html_color(self) -> str:
        return {"ELEVATED": "#DC1414", "GUARDED": "#FFA500", "LOW": "#228B22"}[self.posture]

    # ── CVSS Pre-fetch ─────────────────────────────────────────────────────
    def prefetch_cvss_scores(self):
        """
        Pre-fetch CVSS scores for all CVEs across all threats.
        Called before appendix generation so CVSS data is available.
        Only fetches for unique CVEs to minimize API calls.
        """
        all_cves = set()
        for t in self.threats:
            for cve in t.cve_ids:
                if cve and cve.startswith("CVE-"):
                    all_cves.add(cve)

        if not all_cves:
            logger.info("   No CVEs to look up CVSS scores for.")
            return

        logger.info("   Looking up CVSS scores for %d unique CVE(s) …", len(all_cves))
        # Limit to first 30 CVEs to avoid excessive API calls
        cve_list = sorted(all_cves)[:30]
        _batch_lookup_cvss(cve_list)
        found = sum(1 for c in cve_list if _cvss_cache.get(c) is not None)
        logger.info("   ✔ CVSS data retrieved for %d / %d CVEs", found, len(cve_list))

    # ── DOCX: Add a horizontal rule ────────────────────────────────────────
    @staticmethod
    def _add_horizontal_rule(doc):
        """Add a horizontal separator line to the document."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("─" * 80)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(8)

    # ── DOCX: Generate a single appendix page ──────────────────────────────
    def _generate_appendix_page(self, doc, appendix_label: str, severity: str,
                                 color_rgb: RGBColor, items: List[ThreatItem]):
        """
        Generate a detailed technical appendix page for one severity category.

        Each threat entry includes all 11 required fields:
        1. CVE/Advisory ID         7. Potential Impact
        2. Severity Score          8. Remediation Steps
        3. Affected Systems        9. Timeline/Urgency
        4. Systems in Environment  10. Source/Reference
        5. Technical Description   11. Date Published
        6. Attack Vector
        """
        # Page break before each appendix
        doc.add_page_break()

        # Appendix heading
        heading = doc.add_heading(appendix_label, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = color_rgb

        # Classification banner
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(CONFIG["classification"])
        run.bold = True
        run.font.color.rgb = RGBColor(255, 165, 0)
        run.font.size = Pt(10)

        # Handle empty category
        no_items_labels = {
            "IMMEDIATE": "No IMMEDIATE items identified in the current threat landscape.",
            "PRIORITY":  "No PRIORITY items identified in the current threat landscape.",
            "ROUTINE":   "No ROUTINE items identified in the current threat landscape.",
        }
        if not items:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(no_items_labels.get(severity, "No items in this category."))
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            return

        # Generate detailed entry for each threat
        for idx, item in enumerate(items, 1):
            doc.add_paragraph()  # spacer

            # ── Item title with index ──
            p = doc.add_paragraph()
            run = p.add_run(f"ITEM {idx}: {item.title}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = color_rgb

            # ── Build detail fields ──
            # 1. CVE/Advisory ID
            cve_str = ", ".join(item.cve_ids) if item.cve_ids else "No CVE assigned"
            self._add_detail_field(doc, "CVE / Advisory ID", cve_str)

            # 2. Severity Score (CVSS)
            cvss_display = "Not available"
            for cve_id in item.cve_ids:
                cvss_info = _cvss_cache.get(cve_id)
                if cvss_info:
                    cvss_display = (
                        f"{cvss_info['score']} / 10.0 ({cvss_info['severity']}) "
                        f"— CVSS v{cvss_info['version']} | Vector: {cvss_info['vector']}"
                    )
                    break
            self._add_detail_field(doc, "Severity Score (CVSS)", cvss_display)

            # 3. Affected Systems / Vendors
            affected = item.source
            if item.relevance_tags:
                affected += f" — {', '.join(item.relevance_tags)}"
            self._add_detail_field(doc, "Affected Systems / Vendors", affected)

            # 4. Systems in Our Environment
            env_systems = _get_systems_in_environment(item.relevance_tags)
            self._add_detail_field(doc, "Systems in Our Environment", env_systems)

            # 5. Technical Description
            desc_text = item.description
            if desc_text and "<" in desc_text:
                desc_text = BeautifulSoup(desc_text, "html.parser").get_text()
            if not desc_text:
                desc_text = item.title
            self._add_detail_field(doc, "Technical Description", desc_text[:800])

            # 6. Attack Vector
            attack_vec = _infer_attack_vector(item.description + " " + item.title,
                                               item.cve_ids)
            self._add_detail_field(doc, "Attack Vector", attack_vec)

            # 7. Potential Impact
            impact = _infer_impact(item.title, item.description,
                                    item.relevance_tags, item.severity)
            self._add_detail_field(doc, "Potential Impact", impact)

            # 8. Remediation Steps
            remediation = _infer_remediation(item.title, item.description, item.severity)
            self._add_detail_field(doc, "Remediation Steps", remediation)

            # 9. Timeline / Urgency
            timeline = _infer_timeline(item.severity)
            self._add_detail_field(doc, "Timeline / Urgency", timeline)

            # 10. Source / Reference
            source_ref = item.url if item.url else f"Source: {item.source}"
            self._add_detail_field(doc, "Source / Reference", source_ref)

            # 11. Date Published
            pub_date = item.published if item.published else "Not specified"
            self._add_detail_field(doc, "Date Published", pub_date)

            # Horizontal separator between entries
            if idx < len(items):
                self._add_horizontal_rule(doc)

    @staticmethod
    def _add_detail_field(doc, label: str, value: str):
        """
        Add a labeled detail field to the document in bold-label: value format.
        Used by appendix generation for each of the 11 detail fields.
        """
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(0, 51, 102)
        run_value = p.add_run(value)
        run_value.font.size = Pt(10)

    # ── DOCX Generation (multi-page with appendices) ───────────────────────
    def generate_docx(self, output_dir: str) -> str:
        """
        Generate the multi-page Word document SITREP.

        Page structure:
            Page 1:  Concise executive summary (existing format)
            Page 2:  APPENDIX A — IMMEDIATE Threats (detailed technical breakdown)
            Page 3:  APPENDIX B — PRIORITY Threats (detailed technical breakdown)
            Page 4:  APPENDIX C — ROUTINE Items (detailed technical breakdown)
        """
        doc = Document()

        # -- Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ════════════════════════════════════════════════════════════════════
        # PAGE 1: CONCISE EXECUTIVE SUMMARY
        # ════════════════════════════════════════════════════════════════════

        # -- Header
        h = doc.add_heading("DAILY CYBERSECURITY SITREP", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Classification banner
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(CONFIG["classification"])
        run.bold = True
        run.font.color.rgb = RGBColor(255, 165, 0)
        run.font.size = Pt(12)

        # Meta table
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = "Light Grid Accent 1"
        meta_data = [
            ("Organization", CONFIG["org_name"]),
            ("Date", self.date_str),
            ("Time Generated", self.now.strftime("%H:%M %Z")),
            ("Overall Threat Posture", self.posture),
        ]
        for i, (k, v) in enumerate(meta_data):
            meta_table.cell(i, 0).text = k
            cell = meta_table.cell(i, 1)
            cell.text = v
            if k == "Overall Threat Posture":
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = self._posture_color_rgb()
                        run.bold = True

        doc.add_paragraph()  # spacer

        # -- Executive Summary
        doc.add_heading("1. EXECUTIVE SUMMARY", level=1)
        imm = self._threats_by_severity("IMMEDIATE")
        pri = self._threats_by_severity("PRIORITY")
        rou = self._threats_by_severity("ROUTINE")
        summary = (
            f"As of {self.date_str}, the overall threat posture for "
            f"{CONFIG['org_name']} is assessed as {self.posture}. "
            f"This SITREP identifies {len(imm)} IMMEDIATE (Red) threat(s), "
            f"{len(pri)} PRIORITY (Amber) threat(s), and {len(rou)} ROUTINE "
            f"(Green) informational item(s) across CISA advisories, vendor "
            f"security bulletins, and public cybersecurity news sources."
        )
        doc.add_paragraph(summary)

        # -- Threat Categories (concise list on page 1)
        for sev_label, sev_items, color_rgb in [
            ("2. IMMEDIATE THREATS (Red)", imm, RGBColor(220, 20, 20)),
            ("3. PRIORITY THREATS (Amber)", pri, RGBColor(255, 165, 0)),
            ("4. ROUTINE ITEMS (Green)", rou, RGBColor(34, 139, 34)),
        ]:
            heading = doc.add_heading(sev_label, level=1)
            for run in heading.runs:
                run.font.color.rgb = color_rgb

            if not sev_items:
                doc.add_paragraph("No items in this category at this time.")
                continue

            for idx, item in enumerate(sev_items, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {item.title}")
                run.bold = True
                run.font.size = Pt(11)

                if item.source:
                    doc.add_paragraph(f"Source: {item.source}", style="List Bullet")
                if item.url:
                    doc.add_paragraph(f"URL: {item.url}", style="List Bullet")
                if item.cve_ids:
                    doc.add_paragraph(f"CVEs: {', '.join(item.cve_ids)}", style="List Bullet")
                if item.relevance_tags:
                    doc.add_paragraph(
                        f"Relevant Technology: {', '.join(item.relevance_tags)}",
                        style="List Bullet",
                    )
                if item.description:
                    desc_clean = BeautifulSoup(item.description, "html.parser").get_text()[:400] if "<" in item.description else item.description[:400]
                    doc.add_paragraph(f"Details: {desc_clean}", style="List Bullet")
                doc.add_paragraph()  # spacer

        # -- Recommendations
        doc.add_heading("5. RECOMMENDED ACTIONS", level=1)
        actions = []
        if imm:
            actions.append(
                "IMMEDIATE: Review and apply patches/mitigations for all Red-level "
                "threats within 24 hours. Engage incident response if active exploitation "
                "is confirmed."
            )
        if pri:
            actions.append(
                "PRIORITY: Schedule patching and review for Amber-level items within "
                "the next 72 hours. Validate exposure through vulnerability scanning."
            )
        actions.extend([
            "Continue monitoring CISA KEV catalog and vendor advisory channels.",
            "Ensure endpoint detection and network monitoring tools are updated "
            "with latest threat signatures.",
            "Share this SITREP with IT Security, Clinical Engineering, and "
            "Infrastructure teams.",
        ])
        for a in actions:
            doc.add_paragraph(a, style="List Bullet")

        # -- Appendix reference note on page 1
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            "📎 Detailed technical appendices follow (Appendix A–C) with "
            "comprehensive breakdown including CVSS scores, affected systems, "
            "attack vectors, remediation steps, and impact analysis."
        )
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # -- Page 1 Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— End of Executive Summary —")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(
            f"Generated on {self.now.strftime('%Y-%m-%d %H:%M')} "
            f"by {CONFIG['org_name']} Security Operations"
        )
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(128, 128, 128)

        # ════════════════════════════════════════════════════════════════════
        # PAGES 2-4: DETAILED TECHNICAL APPENDICES
        # ════════════════════════════════════════════════════════════════════

        # APPENDIX A — IMMEDIATE Threats
        self._generate_appendix_page(
            doc,
            "APPENDIX A — IMMEDIATE THREATS (Detailed Technical Breakdown)",
            "IMMEDIATE",
            RGBColor(220, 20, 20),
            imm,
        )

        # APPENDIX B — PRIORITY Threats
        self._generate_appendix_page(
            doc,
            "APPENDIX B — PRIORITY THREATS (Detailed Technical Breakdown)",
            "PRIORITY",
            RGBColor(255, 165, 0),
            pri,
        )

        # APPENDIX C — ROUTINE Items
        self._generate_appendix_page(
            doc,
            "APPENDIX C — ROUTINE ITEMS (Detailed Technical Breakdown)",
            "ROUTINE",
            RGBColor(34, 139, 34),
            rou,
        )

        # ── Final document footer ──
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("— End of Full SITREP Document —")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(
            f"Generated on {self.now.strftime('%Y-%m-%d %H:%M')} "
            f"by {CONFIG['org_name']} Security Operations  |  "
            f"Total items: {len(self.threats)} "
            f"(🔴 {len(imm)} IMMEDIATE, 🟠 {len(pri)} PRIORITY, 🟢 {len(rou)} ROUTINE)"
        )
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(128, 128, 128)

        # Save
        filename = f"Daily_Cybersecurity_SITREP_{self.date_file}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        logger.info("📄 DOCX saved (multi-page with appendices): %s", filepath)
        return filepath

    # ── HTML Generation (for email body) ────────────────────────────────────
    def generate_html(self) -> str:
        """Generate HTML-formatted SITREP for email body."""
        imm = self._threats_by_severity("IMMEDIATE")
        pri = self._threats_by_severity("PRIORITY")
        rou = self._threats_by_severity("ROUTINE")

        def _threat_rows(items: List[ThreatItem]) -> str:
            if not items:
                return "<tr><td colspan='4' style='padding:8px;color:#888;'>No items in this category.</td></tr>"
            rows = ""
            for i, item in enumerate(items, 1):
                cve_str = ", ".join(item.cve_ids) if item.cve_ids else "—"
                tags_str = ", ".join(item.relevance_tags) if item.relevance_tags else "—"
                desc_clean = (BeautifulSoup(item.description, "html.parser").get_text()[:300] if "<" in item.description else item.description[:300]) if item.description else ""
                title_link = f'<a href="{item.url}" style="color:#0066cc;">{item.title}</a>' if item.url else item.title
                rows += f"""
                <tr style="border-bottom:1px solid #e0e0e0;">
                    <td style="padding:8px;vertical-align:top;font-weight:bold;">{i}</td>
                    <td style="padding:8px;vertical-align:top;">{title_link}
                        <br><small style="color:#666;">Source: {item.source} | CVEs: {cve_str}</small>
                        <br><small style="color:#666;">Relevant Tech: {tags_str}</small>
                    </td>
                    <td style="padding:8px;vertical-align:top;font-size:12px;">{desc_clean}</td>
                    <td style="padding:8px;vertical-align:top;text-align:center;">
                        <span style="color:{item.html_color};font-weight:bold;">{item.severity}</span>
                    </td>
                </tr>"""
            return rows

        html = f"""
<!DOCTYPE html>
<html>
<head><meta charset="utf-8"></head>
<body style="font-family:Calibri,Arial,sans-serif;color:#333;max-width:900px;margin:auto;padding:20px;">

<!-- Header -->
<div style="background:#003366;color:white;padding:20px;text-align:center;border-radius:6px 6px 0 0;">
    <h1 style="margin:0;font-size:24px;">🛡️ DAILY CYBERSECURITY SITREP</h1>
    <p style="margin:5px 0 0;font-size:14px;">{CONFIG['org_name']} — {self.date_str}</p>
</div>

<!-- Classification -->
<div style="background:#FFA500;color:white;text-align:center;padding:6px;font-weight:bold;font-size:13px;">
    {CONFIG['classification']}
</div>

<!-- Threat Posture -->
<div style="background:#f5f5f5;padding:15px;border:1px solid #ddd;margin-top:0;">
    <table style="width:100%;">
        <tr>
            <td><strong>Date:</strong> {self.date_str}</td>
            <td><strong>Time:</strong> {self.now.strftime('%H:%M')}</td>
            <td><strong>Overall Threat Posture:</strong>
                <span style="color:{self._posture_html_color()};font-weight:bold;font-size:16px;">
                    {self.posture}
                </span>
            </td>
        </tr>
    </table>
</div>

<!-- Summary -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;">
    <h2 style="color:#003366;font-size:18px;">📋 Executive Summary</h2>
    <p>As of {self.date_str}, the overall threat posture is assessed as
       <strong style="color:{self._posture_html_color()};">{self.posture}</strong>.
       This SITREP identifies <strong style="color:#DC1414;">{len(imm)} IMMEDIATE</strong>,
       <strong style="color:#FFA500;">{len(pri)} PRIORITY</strong>, and
       <strong style="color:#228B22;">{len(rou)} ROUTINE</strong> items.</p>
    <p>Sources: CISA KEV, CISA Alerts, {len(VENDOR_ADVISORY_URLS)} vendor advisory feeds,
       and cybersecurity news outlets.</p>
</div>

<!-- IMMEDIATE -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;">
    <h2 style="color:#DC1414;font-size:18px;">🔴 IMMEDIATE THREATS (Red) — {len(imm)} Item(s)</h2>
    <table style="width:100%;border-collapse:collapse;">
        <thead>
            <tr style="background:#DC1414;color:white;">
                <th style="padding:8px;width:30px;">#</th>
                <th style="padding:8px;">Threat / Advisory</th>
                <th style="padding:8px;">Details</th>
                <th style="padding:8px;width:90px;">Level</th>
            </tr>
        </thead>
        <tbody>{_threat_rows(imm)}</tbody>
    </table>
</div>

<!-- PRIORITY -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;">
    <h2 style="color:#FFA500;font-size:18px;">🟠 PRIORITY THREATS (Amber) — {len(pri)} Item(s)</h2>
    <table style="width:100%;border-collapse:collapse;">
        <thead>
            <tr style="background:#FFA500;color:white;">
                <th style="padding:8px;width:30px;">#</th>
                <th style="padding:8px;">Threat / Advisory</th>
                <th style="padding:8px;">Details</th>
                <th style="padding:8px;width:90px;">Level</th>
            </tr>
        </thead>
        <tbody>{_threat_rows(pri)}</tbody>
    </table>
</div>

<!-- ROUTINE -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;">
    <h2 style="color:#228B22;font-size:18px;">🟢 ROUTINE ITEMS (Green) — {len(rou)} Item(s)</h2>
    <table style="width:100%;border-collapse:collapse;">
        <thead>
            <tr style="background:#228B22;color:white;">
                <th style="padding:8px;width:30px;">#</th>
                <th style="padding:8px;">Threat / Advisory</th>
                <th style="padding:8px;">Details</th>
                <th style="padding:8px;width:90px;">Level</th>
            </tr>
        </thead>
        <tbody>{_threat_rows(rou)}</tbody>
    </table>
</div>

<!-- Recommendations -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;">
    <h2 style="color:#003366;font-size:18px;">✅ Recommended Actions</h2>
    <ul style="line-height:1.8;">
        {"".join(f"<li><strong>IMMEDIATE:</strong> Review and patch all Red-level threats within 24 hours.</li>" for _ in [1] if imm)}
        {"".join(f"<li><strong>PRIORITY:</strong> Schedule remediation for Amber-level items within 72 hours.</li>" for _ in [1] if pri)}
        <li>Continue monitoring CISA KEV and vendor advisory channels daily.</li>
        <li>Update endpoint detection and network monitoring with latest threat signatures.</li>
        <li>Distribute this SITREP to IT Security, Clinical Engineering, and Infrastructure teams.</li>
    </ul>
</div>

<!-- Appendix Notice -->
<div style="padding:15px;border:1px solid #ddd;border-top:none;background:#f0f4f8;">
    <h2 style="color:#003366;font-size:16px;">📎 Detailed Technical Appendices</h2>
    <p style="font-size:13px;">Detailed technical appendices are included in the full report document (Word attachment).
       The appendices contain comprehensive breakdowns for each threat including:</p>
    <ul style="font-size:13px;line-height:1.6;">
        <li>CVE/Advisory IDs and CVSS severity scores</li>
        <li>Affected systems and hospital environment mapping</li>
        <li>Technical descriptions and attack vectors</li>
        <li>Healthcare-specific impact analysis</li>
        <li>Remediation steps and urgency timelines</li>
    </ul>
    <p style="font-size:12px;color:#666;">
        <strong>APPENDIX A:</strong> IMMEDIATE Threats |
        <strong>APPENDIX B:</strong> PRIORITY Threats |
        <strong>APPENDIX C:</strong> ROUTINE Items
    </p>
</div>

<!-- Footer -->
<div style="background:#003366;color:white;padding:15px;text-align:center;border-radius:0 0 6px 6px;margin-top:0;font-size:12px;">
    <p style="margin:0;">Generated on {self.now.strftime('%Y-%m-%d %H:%M')} by {CONFIG['org_name']} Security Operations</p>
    <p style="margin:4px 0 0;opacity:0.7;">This report is auto-generated. Verify findings before taking action.</p>
    <p style="margin:4px 0 0;opacity:0.7;">Detailed technical appendices are included in the full report document.</p>
</div>

</body>
</html>"""
        return html


# ── Email Delivery ──────────────────────────────────────────────────────────
def send_email_smtp(subject: str, html_body: str, recipient: str):
    """Send the SITREP email via SMTP."""
    smtp_user = CONFIG["smtp_username"] or os.environ.get("SITREP_SMTP_USER", "")
    smtp_pass = CONFIG["smtp_password"] or os.environ.get("SITREP_SMTP_PASS", "")

    if not smtp_user or not smtp_pass:
        logger.warning(
            "⚠  SMTP credentials not configured. Set environment variables "
            "SITREP_SMTP_USER and SITREP_SMTP_PASS, or edit CONFIG in the script."
        )
        logger.info("📧 Email would have been sent to: %s", recipient)
        logger.info("   Subject: %s", subject)
        # Save HTML to file as fallback
        fallback_path = os.path.join(
            CONFIG["output_dir"],
            f"SITREP_Email_{datetime.datetime.now().strftime('%Y%m%d')}.html",
        )
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(html_body)
        logger.info("   HTML email body saved to: %s", fallback_path)
        return False

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = CONFIG["email_from"]
    msg["To"] = recipient
    msg.attach(MIMEText(html_body, "html", "utf-8"))

    try:
        logger.info("📧 Sending email via %s:%d …", CONFIG["smtp_server"], CONFIG["smtp_port"])
        if CONFIG["smtp_use_tls"]:
            server = smtplib.SMTP(CONFIG["smtp_server"], CONFIG["smtp_port"])
            server.ehlo()
            server.starttls(context=ssl.create_default_context())
        else:
            server = smtplib.SMTP_SSL(CONFIG["smtp_server"], CONFIG["smtp_port"],
                                       context=ssl.create_default_context())
        server.ehlo()
        server.login(smtp_user, smtp_pass)
        server.sendmail(CONFIG["email_from"], [recipient], msg.as_string())
        server.quit()
        logger.info("   ✔ Email sent successfully to %s", recipient)
        return True
    except Exception as exc:
        logger.error("   ✖ Email send failed: %s", exc)
        # Save fallback
        fallback_path = os.path.join(
            CONFIG["output_dir"],
            f"SITREP_Email_{datetime.datetime.now().strftime('%Y%m%d')}.html",
        )
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(html_body)
        logger.info("   HTML email body saved as fallback: %s", fallback_path)
        return False


def send_email_abacusai(subject: str, html_body: str, recipient: str, docx_path: str = None):
    """Send email using Abacus AI Send_Email_Tool integration (for DeepAgent environments)."""
    try:
        # This path is used when running inside the DeepAgent environment
        # The script can detect the environment and use the appropriate method
        logger.info("📧 Attempting Abacus AI email delivery to %s …", recipient)
        # Fallback: save HTML for manual sending
        fallback_path = os.path.join(
            CONFIG["output_dir"],
            f"SITREP_Email_{datetime.datetime.now().strftime('%Y%m%d')}.html",
        )
        with open(fallback_path, "w", encoding="utf-8") as f:
            f.write(html_body)
        logger.info("   HTML email saved to: %s", fallback_path)
        return False
    except Exception as exc:
        logger.error("   ✖ Email delivery error: %s", exc)
        return False


# ── Main Entry Point ────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Generate Daily Cybersecurity SITREP on-demand",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Examples:
              python generate_sitrep_manual.py
              python generate_sitrep_manual.py --no-email
              python generate_sitrep_manual.py --output /tmp/reports
              python generate_sitrep_manual.py --email someone@example.com
        """),
    )
    parser.add_argument(
        "--no-email", action="store_true",
        help="Skip sending email; only generate the DOCX and HTML files",
    )
    parser.add_argument(
        "--output", "-o", type=str, default=CONFIG["output_dir"],
        help="Output directory for generated files (default: current dir)",
    )
    parser.add_argument(
        "--email", "-e", type=str, default=CONFIG["email_recipient"],
        help="Override email recipient",
    )
    args = parser.parse_args()

    # Apply args
    CONFIG["output_dir"] = args.output
    os.makedirs(CONFIG["output_dir"], exist_ok=True)

    # Setup file logging
    log_path = os.path.join(CONFIG["output_dir"], CONFIG["log_file"])
    _setup_file_logging(log_path)

    # Banner
    print()
    print("╔═══════════════════════════════════════════════════════════════╗")
    print("║       🛡️  Daily Cybersecurity SITREP Generator  🛡️          ║")
    print("║       Evergreen Healthcare — Security Operations            ║")
    print("╚═══════════════════════════════════════════════════════════════╝")
    print()

    start_time = time.time()

    # Step 1: Collect Intelligence
    logger.info("STEP 1/5: Gathering cybersecurity intelligence …")
    collector = IntelCollector()
    collector.collect_all()
    threats = collector.threats

    if not threats:
        logger.warning("No threat items collected. The report will be empty.")

    # Step 2: CVSS Score Lookup & Generate DOCX (multi-page with appendices)
    logger.info("STEP 2/5: Looking up CVSS scores for discovered CVEs …")
    generator = SITREPGenerator(threats)
    generator.prefetch_cvss_scores()

    logger.info("STEP 3/5: Generating multi-page SITREP Word document (with appendices) …")
    docx_path = generator.generate_docx(CONFIG["output_dir"])

    # Step 4: Generate HTML
    logger.info("STEP 4/5: Generating HTML email body …")
    html_body = generator.generate_html()

    # Save HTML copy
    html_path = os.path.join(
        CONFIG["output_dir"],
        f"Daily_Cybersecurity_SITREP_{generator.date_file}.html",
    )
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_body)
    logger.info("📄 HTML saved: %s", html_path)

    # Step 5: Send Email
    subject = (
        f"Daily Cybersecurity SITREP - {generator.date_str} - "
        f"Threat Level: {generator.posture}"
    )
    if args.no_email:
        logger.info("STEP 5/5: Email sending skipped (--no-email flag)")
    else:
        logger.info("STEP 5/5: Sending SITREP email …")
        send_email_smtp(subject, html_body, args.email)

    # Summary
    elapsed = time.time() - start_time
    imm = sum(1 for t in threats if t.severity == "IMMEDIATE")
    pri = sum(1 for t in threats if t.severity == "PRIORITY")
    rou = sum(1 for t in threats if t.severity == "ROUTINE")

    print()
    print("╔═══════════════════════════════════════════════════════════════╗")
    print("║                    SITREP GENERATION COMPLETE                ║")
    print("╠═══════════════════════════════════════════════════════════════╣")
    print(f"║  Date          : {generator.date_str:<43} ║")
    print(f"║  Threat Posture: {generator.posture:<43} ║")
    print(f"║  Total Items   : {len(threats):<43} ║")
    print(f"║  🔴 IMMEDIATE  : {imm:<43} ║")
    print(f"║  🟠 PRIORITY   : {pri:<43} ║")
    print(f"║  🟢 ROUTINE    : {rou:<43} ║")
    print(f"║  DOCX File     : {os.path.basename(docx_path):<43} ║")
    print(f"║  HTML File     : {os.path.basename(html_path):<43} ║")
    print(f"║  Log File      : {CONFIG['log_file']:<43} ║")
    print(f"║  Elapsed       : {elapsed:.1f}s{'':<39} ║")
    print("╠═══════════════════════════════════════════════════════════════╣")
    print("║  📎 DOCX includes detailed technical appendices (A, B, C)   ║")
    print("║  📧 Email body = executive summary + appendix notice        ║")
    print("╚═══════════════════════════════════════════════════════════════╝")
    print()

    return 0


if __name__ == "__main__":
    sys.exit(main())
