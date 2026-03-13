from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document('Daily_Cybersecurity_SITREP_Template.docx')

# Color constants matching existing doc
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x57)
GREY_TEXT = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GREY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xD4, 0x8B, 0x0B)
GREEN = RGBColor(0x27, 0x7A, 0x3E)

def add_run(paragraph, text, size=Pt(9.5), bold=False, color=GREY_TEXT, font_name='Calibri'):
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add bottom border to paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def add_threat_entry(doc, number, accent_color):
    # Threat title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f'Threat #{number}: ', size=Pt(11), bold=True, color=accent_color)
    add_run(p, '[Threat Title/Name]', size=Pt(11), bold=True, color=GREY_TEXT)

    fields = [
        ('CVE/Advisory ID:', '[CVE-YYYY-XXXXX or Vendor Advisory Number]'),
        ('Severity Score:', '[CVSS Score if available, e.g., 9.8 Critical]'),
        ('Affected Systems/Vendors:', '[Specific products, versions, e.g., Vendor Product v1.2.x – v1.4.x]'),
        ('Systems in Our Environment:', '[Which of our hospital systems are affected, e.g., 12 servers in Radiology VLAN]'),
        ('Technical Description:', '[Detailed technical explanation of the vulnerability or threat — describe the flaw, affected component, and how it manifests]'),
        ('Attack Vector:', '[How the threat could be exploited, e.g., Network/Remote, requires no authentication, exploitable via crafted HTTP request]'),
        ('Potential Impact:', '[Detailed impact analysis specific to healthcare/hospital operations — patient safety, PHI exposure, system downtime, regulatory implications]'),
        ('Remediation Steps:', '[Specific technical actions required — patches, configuration changes, workarounds, compensating controls]'),
        ('Timeline/Urgency:', '[When action must be taken, e.g., Patch within 24 hours / Apply workaround immediately / Monitor — next patch cycle]'),
        ('Source/Reference:', '[URL or source of the advisory, e.g., https://www.cisa.gov/known-exploited-vulnerabilities-catalog]'),
        ('Date Published:', '[When the advisory was released, e.g., MM/DD/YYYY]'),
    ]

    for label, placeholder in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.25)
        add_run(p, f'  {label} ', size=Pt(9), bold=True, color=DARK_BLUE)
        add_run(p, placeholder, size=Pt(9), bold=False, color=LIGHT_GREY)


def add_appendix_page(doc, letter, title, accent_color):
    import docx.enum.text as wet
    
    # Page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(wet.WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Colored accent bar (top border on header paragraph)
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(4)
    # Add top border as accent
    pPr = header_p._element.get_or_add_pPr()
    color_hex = str(accent_color)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="24" w:space="4" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)
    add_run(header_p, f'APPENDIX {letter}: {title}', size=Pt(14), bold=True, color=accent_color)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(2)
    add_run(sub_p, 'DETAILED TECHNICAL BREAKDOWN', size=Pt(10), bold=True, color=DARK_BLUE)

    # Note box
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(8)
    # Add shading to note
    pPr = note_p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
    pPr.append(shd)
    # Add border around note
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                     f'<w:left w:val="single" w:sz="12" w:space="8" w:color="{color_hex}"/>'
                     f'<w:top w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>'
                     f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>'
                     f'<w:right w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)
    note_p.paragraph_format.left_indent = Inches(0.1)
    note_p.paragraph_format.right_indent = Inches(0.1)
    add_run(note_p, 'ℹ ', size=Pt(9), bold=True, color=accent_color)
    add_run(note_p, 'This appendix provides detailed technical information for items summarized on Page 1. Intended for technical staff and incident response teams.', 
            size=Pt(9), bold=False, color=GREY_TEXT)

    # Threat entry 1
    add_threat_entry(doc, 1, accent_color)

    # Horizontal line separator
    add_horizontal_line(doc)

    # Threat entry 2
    add_threat_entry(doc, 2, accent_color)

    # Bottom note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, '[Add additional threat entries as needed following the same format above]', size=Pt(8.5), bold=False, color=LIGHT_GREY)


# Add the three appendix pages
add_appendix_page(doc, 'A', 'IMMEDIATE THREATS', RED)
add_appendix_page(doc, 'B', 'PRIORITY THREATS', AMBER)
add_appendix_page(doc, 'C', 'ROUTINE ITEMS', GREEN)

doc.save('Daily_Cybersecurity_SITREP_Template.docx')
print('Done - saved successfully')
