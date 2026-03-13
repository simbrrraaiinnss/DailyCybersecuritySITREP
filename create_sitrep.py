from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# Colors
NAVY = RGBColor(0x1B, 0x3A, 0x57)
DARK_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
AMBER_ACCENT = RGBColor(0xD4, 0x8B, 0x0B)
GREEN_ACCENT = RGBColor(0x27, 0x7A, 0x3E)
NAVY_HEX = "1B3A57"
LIGHT_BG = "F0F3F6"
RED_HEX = "C0392B"
AMBER_HEX = "D48B0B"
GREEN_HEX = "277A3E"
TABLE_HEADER_BG = "1B3A57"
TABLE_ALT_BG = "F5F7FA"
BORDER_COLOR = "B0BEC5"

doc = Document()

# ── Global style defaults ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)
font.color.rgb = DARK_GRAY
style.paragraph_format.space_before = Pt(1)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.08

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

# ── Helper functions ──
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", BORDER_COLOR)}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc_or_cell, text, font_size=9.5, bold=False, color=DARK_GRAY,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2,
                             italic=False, font_name='Calibri'):
    p = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, 'add_paragraph') else doc_or_cell.paragraphs[0] if not text else doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    run.font.italic = italic
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    return p

def add_section_heading(doc, text, color, underline_hex=None, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    # Add bottom border
    hex_color = underline_hex or f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{hex_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, color=DARK_GRAY, bold_prefix="", font_size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.size = Pt(font_size)
        run_b.font.bold = True
        run_b.font.color.rgb = color
        run_b.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.left_indent = Cm(0.8)
    return p

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def format_table_cell(cell, text, font_size=8.5, bold=False, color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def remove_table_borders(table):
    """Set clean thin borders on entire table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# ══════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════
header = doc.sections[0].header
# Clear default
for p in header.paragraphs:
    p.clear()

# Header table: logo | title block
htable = header.add_table(rows=1, cols=2, width=Inches(6.8))
htable.alignment = WD_TABLE_ALIGNMENT.CENTER
htable.columns[0].width = Inches(1.5)
htable.columns[1].width = Inches(5.3)

# Logo placeholder cell
logo_cell = htable.cell(0, 0)
logo_cell.text = ""
lp = logo_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Draw a placeholder box
run = lp.add_run("[Hospital Logo]")
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GRAY
run.font.name = 'Calibri'
run.font.italic = True
lp.paragraph_format.space_before = Pt(6)
lp.paragraph_format.space_after = Pt(2)
set_cell_shading(logo_cell, LIGHT_BG)
logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Title cell
title_cell = htable.cell(0, 1)
title_cell.text = ""
tp1 = title_cell.paragraphs[0]
r1 = tp1.add_run("[HOSPITAL NAME]")
r1.font.size = Pt(16)
r1.font.bold = True
r1.font.color.rgb = NAVY
r1.font.name = 'Calibri'
tp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tp1.paragraph_format.space_before = Pt(2)
tp1.paragraph_format.space_after = Pt(0)

tp2 = title_cell.add_paragraph()
r2 = tp2.add_run("Information Security Department")
r2.font.size = Pt(10)
r2.font.color.rgb = LIGHT_GRAY
r2.font.name = 'Calibri'
tp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tp2.paragraph_format.space_before = Pt(0)
tp2.paragraph_format.space_after = Pt(2)
title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Remove header table borders
tbl = htable._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{NAVY_HEX}"/>'
    f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

# ══════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════
footer = doc.sections[0].footer
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.clear()

# Add top border to footer paragraph
fpPr = fp._p.get_or_add_pPr()
fpBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="1" w:color="{NAVY_HEX}"/>'
    f'</w:pBdr>'
)
fpPr.append(fpBdr)

# Use tab stops for left/right alignment
tab_stops = fp.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(6.8), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

run_conf = fp.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run_conf.font.size = Pt(7.5)
run_conf.font.bold = True
run_conf.font.color.rgb = NAVY
run_conf.font.name = 'Calibri'

fp.add_run("\t")

# Page number field
run_page = fp.add_run()
run_page.font.size = Pt(7.5)
run_page.font.color.rgb = LIGHT_GRAY
run_page.font.name = 'Calibri'
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_page._r.append(fldChar1)
run_page2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run_page2._r.append(instrText)
run_page3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_page3._r.append(fldChar2)

fp.paragraph_format.space_before = Pt(4)
fp.paragraph_format.space_after = Pt(0)

# ══════════════════════════════════════════════════════════════
# DOCUMENT TITLE
# ══════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_run = title_p.add_run("DAILY CYBERSECURITY SITUATIONAL REPORT (SITREP)")
title_run.font.size = Pt(13)
title_run.font.bold = True
title_run.font.color.rgb = NAVY
title_run.font.name = 'Calibri'
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(2)
title_p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════
# 1. REPORT HEADER METADATA TABLE
# ══════════════════════════════════════════════════════════════
meta_table = doc.add_table(rows=3, cols=4)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.autofit = True

meta_data = [
    ["Report Date:", "[MM/DD/YYYY]", "Report Number:", "[YYYYMMDD-SITREP]"],
    ["Prepared By:", "[Analyst Name]", "Classification:", "[CONFIDENTIAL / INTERNAL]"],
]

for i, row_data in enumerate(meta_data):
    for j, text in enumerate(row_data):
        cell = meta_table.cell(i, j)
        is_label = j % 2 == 0
        format_table_cell(cell, text, font_size=9, bold=is_label, 
                         color=NAVY if is_label else DARK_GRAY)
        if is_label:
            set_cell_shading(cell, LIGHT_BG)

# Threat posture row - merge all cells
threat_cell_start = meta_table.cell(2, 0)
threat_cell_end = meta_table.cell(2, 3)
merged = threat_cell_start.merge(threat_cell_end)
merged.text = ""
tp = merged.paragraphs[0]
r1 = tp.add_run("Overall Threat Posture:    ")
r1.font.size = Pt(10)
r1.font.bold = True
r1.font.color.rgb = NAVY
r1.font.name = 'Calibri'
# Threat level badges
for level, color_hex in [("ELEVATED", RED_HEX), ("GUARDED", AMBER_HEX), ("LOW", GREEN_HEX)]:
    r = tp.add_run(f"  {level}  ")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = LIGHT_GRAY
    r.font.name = 'Calibri'
    if level == "ELEVATED":
        r.font.color.rgb = LIGHT_GRAY
    r2 = tp.add_run("  /  ") if level != "LOW" else None
    if r2:
        r2.font.size = Pt(9)
        r2.font.color.rgb = LIGHT_GRAY

tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.space_before = Pt(3)
tp.paragraph_format.space_after = Pt(3)
set_cell_shading(merged, NAVY_HEX)
# Override text colors for threat row
for run in tp.runs:
    if run.text.strip() in ["ELEVATED", "GUARDED", "LOW", "/"]:
        pass
# Redo the threat posture row properly
merged.text = ""
tp = merged.paragraphs[0]
r1 = tp.add_run("  OVERALL THREAT POSTURE:     ")
r1.font.size = Pt(10)
r1.font.bold = True
r1.font.color.rgb = WHITE
r1.font.name = 'Calibri'

# Circle one instruction
r_inst = tp.add_run("[ ELEVATED  /  GUARDED  /  LOW ]")
r_inst.font.size = Pt(10)
r_inst.font.bold = True
r_inst.font.color.rgb = RGBColor(0xA0, 0xC0, 0xD8)
r_inst.font.name = 'Calibri'

tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.space_before = Pt(4)
tp.paragraph_format.space_after = Pt(4)

remove_table_borders(meta_table)
# Override with navy border for meta table
tbl = meta_table._tbl
for existing in tbl.tblPr.findall(qn('w:tblBorders')):
    tbl.tblPr.remove(existing)
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{NAVY_HEX}"/>'
    f'<w:left w:val="single" w:sz="6" w:space="0" w:color="{NAVY_HEX}"/>'
    f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{NAVY_HEX}"/>'
    f'<w:right w:val="single" w:sz="6" w:space="0" w:color="{NAVY_HEX}"/>'
    f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
    f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
    f'</w:tblBorders>'
)
tbl.tblPr.append(borders)

# ══════════════════════════════════════════════════════════════
# 2. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "EXECUTIVE SUMMARY", NAVY)
add_formatted_paragraph(doc, 
    "[Provide a 2–3 sentence plain-language overview of the current threat landscape as it relates to "
    "healthcare. Focus on what the CIO needs to know at a glance — no technical jargon.]",
    font_size=9.5, italic=True, color=LIGHT_GRAY)

# ══════════════════════════════════════════════════════════════
# 3. IMMEDIATE — Action Required
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "⬤  IMMEDIATE — Action Required", RED_ACCENT, RED_HEX)

note_p = doc.add_paragraph()
note_r = note_p.add_run("If no items, write: ")
note_r.font.size = Pt(8)
note_r.font.color.rgb = LIGHT_GRAY
note_r.font.italic = True
note_r.font.name = 'Calibri'
note_r2 = note_p.add_run('"No immediate items to report."')
note_r2.font.size = Pt(8)
note_r2.font.color.rgb = LIGHT_GRAY
note_r2.font.italic = True
note_r2.font.name = 'Calibri'
note_p.paragraph_format.space_before = Pt(0)
note_p.paragraph_format.space_after = Pt(2)

# Action Required Table
action_table = doc.add_table(rows=3, cols=5)
action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
action_table.autofit = True

headers = ["#", "Description", "Source", "Impact to Our Org", "Recommended Action"]
col_widths = [Inches(0.35), Inches(2.0), Inches(1.0), Inches(1.6), Inches(1.85)]

for j, (header_text, width) in enumerate(zip(headers, col_widths)):
    cell = action_table.cell(0, j)
    format_table_cell(cell, header_text, font_size=8, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, RED_HEX)
    action_table.columns[j].width = width

sample_rows = [
    ["1", "Critical RCE vulnerability in [EHR vendor] — active exploitation reported (CVE-XXXX-XXXXX)",
     "H-ISAC Alert", "Direct — we run this system in production",
     "Apply emergency patch within 24 hrs; coordinate with vendor"],
    ["2", "Ransomware campaign targeting healthcare orgs via phishing with .ISO attachments",
     "HHS HC3", "High — similar phishing attempts seen in our email logs",
     "Block .ISO attachments at email gateway; send staff awareness alert"],
]

for i, row_data in enumerate(sample_rows):
    for j, text in enumerate(row_data):
        cell = action_table.cell(i + 1, j)
        format_table_cell(cell, text, font_size=8.5, color=DARK_GRAY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    if i % 2 == 1:
        for j in range(5):
            set_cell_shading(action_table.cell(i + 1, j), TABLE_ALT_BG)

remove_table_borders(action_table)

# ══════════════════════════════════════════════════════════════
# 4. PRIORITY — Awareness & Preparation
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "⬤  PRIORITY — Awareness & Preparation", AMBER_ACCENT, AMBER_HEX)
add_bullet(doc, "CISA adds three new vulnerabilities to the Known Exploited Vulnerabilities catalog affecting network "
           "infrastructure devices — review our asset inventory for exposure. [Source: CISA KEV]", DARK_GRAY)
add_bullet(doc, "Emerging social engineering campaign impersonating healthcare compliance auditors via phone calls "
           "to IT help desks — brief help desk staff on verification procedures. [Source: FBI InfraGard]", DARK_GRAY)

# ══════════════════════════════════════════════════════════════
# 5. ROUTINE — General Awareness
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "⬤  ROUTINE — General Awareness", GREEN_ACCENT, GREEN_HEX)
add_bullet(doc, "HHS releases updated guidance on HIPAA Security Rule compliance for cloud-hosted systems — "
           "review for applicability to our Azure/AWS environments. [Source: HHS]", DARK_GRAY)
add_bullet(doc, "Industry report indicates 45% increase in credential-stuffing attacks against patient portals "
           "sector-wide in Q4 — our MFA implementation mitigates this risk. [Source: H-ISAC Quarterly Report]", DARK_GRAY)

# ══════════════════════════════════════════════════════════════
# 6. WHAT THIS MEANS FOR US
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "WHAT THIS MEANS FOR US", NAVY)
add_formatted_paragraph(doc,
    "[Connect the above items specifically to our hospital's environment, systems, and vendors. "
    "Example: 'The EHR vulnerability in Item #1 directly affects our Epic deployment. We are coordinating "
    "with Epic support for emergency patching tonight during the maintenance window. The phishing campaign "
    "in Item #2 aligns with suspicious emails flagged by our SOC yesterday — additional email filtering "
    "rules have been deployed.']",
    font_size=9.5, italic=True, color=LIGHT_GRAY)

# ══════════════════════════════════════════════════════════════
# 7. ACTIONS TAKEN IN THE LAST 24 HOURS
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "ACTIONS TAKEN IN THE LAST 24 HOURS", NAVY)
add_bullet(doc, "Deployed updated endpoint detection signatures across all workstations and servers.", DARK_GRAY)
add_bullet(doc, "Completed vulnerability scan of DMZ-facing systems — 3 findings under remediation.", DARK_GRAY)
add_bullet(doc, "Coordinated with [Vendor Name] on patch timeline for identified critical vulnerability.", DARK_GRAY)

# ══════════════════════════════════════════════════════════════
# 8. OPEN ITEMS / WATCH LIST
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "OPEN ITEMS / WATCH LIST", NAVY)
add_bullet(doc, "[Day 3] — Awaiting vendor patch for CVE-XXXX-XXXXX affecting radiology imaging system; "
           "compensating controls in place.", DARK_GRAY)
add_bullet(doc, "[Day 7] — Monitoring anomalous outbound traffic from Lab subnet; investigation ongoing, "
           "no confirmed compromise.", DARK_GRAY)

# ══════════════════════════════════════════════════════════════
# 9. UPCOMING DEADLINES & COMPLIANCE ITEMS
# ══════════════════════════════════════════════════════════════
add_section_heading(doc, "UPCOMING DEADLINES & COMPLIANCE ITEMS", NAVY)
add_bullet(doc, "[MM/DD] — Monthly critical patch cycle deadline (HIPAA Security Rule §164.308(a)(5)).", DARK_GRAY)
add_bullet(doc, "[MM/DD] — Annual penetration test report due to Compliance Office.", DARK_GRAY)

# ══════════════════════════════════════════════════════════════
# 10. REPORT FOOTER
# ══════════════════════════════════════════════════════════════
# Thin separator line
sep_p = doc.add_paragraph()
sep_pPr = sep_p._p.get_or_add_pPr()
sep_bdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="1" w:color="{NAVY_HEX}"/>'
    f'</w:pBdr>'
)
sep_pPr.append(sep_bdr)
sep_p.paragraph_format.space_before = Pt(6)
sep_p.paragraph_format.space_after = Pt(2)

footer_table = doc.add_table(rows=1, cols=3)
footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
footer_data = [
    ("Prepared by:", "[Name, Title]"),
    ("Next SITREP:", "[Date/Time]"),
    ("Distribution:", "[CIO, CISO, IT Directors]"),
]
for j, (label, value) in enumerate(footer_data):
    cell = footer_table.cell(0, j)
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(label + " ")
    r1.font.size = Pt(8)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(8)
    r2.font.color.rgb = DARK_GRAY
    r2.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# Remove borders from footer table
tbl = footer_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)


# ══════════════════════════════════════════════════════════════
# PAGE BREAK — USAGE GUIDE (PAGE 2)
# ══════════════════════════════════════════════════════════════
doc.add_page_break()

# Usage guide title
guide_title = doc.add_paragraph()
gt_run = guide_title.add_run("TEMPLATE USAGE GUIDE")
gt_run.font.size = Pt(14)
gt_run.font.bold = True
gt_run.font.color.rgb = NAVY
gt_run.font.name = 'Calibri'
guide_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
guide_title.paragraph_format.space_before = Pt(4)
guide_title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph()
st_run = subtitle.add_run("⚠  Remove this page before distributing the report  ⚠")
st_run.font.size = Pt(10)
st_run.font.bold = True
st_run.font.color.rgb = RED_ACCENT
st_run.font.name = 'Calibri'
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(8)

# Guide sections
def guide_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def guide_text(text, bold=False, italic=False):
    return add_formatted_paragraph(doc, text, font_size=9.5, bold=bold, italic=italic, color=DARK_GRAY)

def guide_bullet(text, bold_prefix=""):
    return add_bullet(doc, text, DARK_GRAY, bold_prefix, font_size=9.5)

# Section: How to Fill Out Each Section
guide_heading("How to Fill Out Each Section")
guide_bullet("Replace all bracketed placeholder text [like this] with actual content.", "Report Header: ")
guide_bullet("Select ONE threat posture level and delete the others. Bold or highlight the selected level.", "Threat Posture: ")
guide_bullet("Write 2–3 sentences summarizing the most important items. A busy CIO should be able to read only this section and understand the day's risk picture.", "Executive Summary: ")
guide_bullet("Only items requiring action within 24–48 hours. If none, state 'No immediate items to report.' Do not leave blank.", "IMMEDIATE (Red): ")
guide_bullet("Emerging threats and developments that require awareness or preparation but not immediate action.", "PRIORITY (Amber): ")
guide_bullet("General industry context, informational advisories, and trends. Helps the CIO stay informed.", "ROUTINE (Green): ")
guide_bullet("This is the most valuable section. Translate technical findings into business impact specific to our hospital's systems, vendors, and patient care operations.", "What This Means for Us: ")
guide_bullet("Document completed security actions to demonstrate operational tempo and accountability.", "Actions Taken: ")
guide_bullet("Track items across multiple days. Include the day count [Day X] so leadership can see how long items have been open.", "Open Items: ")
guide_bullet("Include regulatory deadlines, patching windows, audit dates, and compliance milestones.", "Deadlines: ")

# Section: Tone & Writing Guidance
guide_heading("Tone & Writing Guidance")
guide_bullet("Write for a non-technical executive. Avoid jargon, acronyms (unless widely known), and deep technical detail.")
guide_bullet("Be concise and factual. Use short, declarative sentences.")
guide_bullet('Use plain language: say "attackers are targeting hospital email systems" not "threat actors are leveraging spear-phishing vectors against healthcare MX endpoints."')
guide_bullet("Keep the main report to ONE page (two pages maximum in exceptional circumstances).")
guide_bullet("Every item should answer: What happened? Does it affect us? What are we doing about it?")

# Section: Threat Level Definitions
guide_heading("Threat Posture Definitions")

# Threat level table
tl_table = doc.add_table(rows=4, cols=3)
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tl_headers = ["Level", "Definition", "Typical Triggers"]
for j, h in enumerate(tl_headers):
    cell = tl_table.cell(0, j)
    format_table_cell(cell, h, font_size=9, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, NAVY_HEX)

tl_data = [
    ("ELEVATED", RED_HEX, "Active threats directly targeting healthcare sector or our organization",
     "Active exploitation of healthcare systems; ransomware campaigns targeting hospitals; direct threat intelligence received"),
    ("GUARDED", AMBER_HEX, "Increased risk environment with no direct threats observed against us",
     "New critical vulnerabilities disclosed; sector-wide advisories issued; increased threat actor activity in healthcare"),
    ("LOW", GREEN_HEX, "Normal operations; no significant threats identified",
     "Routine patching cycle; no active campaigns targeting healthcare; stable threat landscape"),
]

for i, (level, color_hex, definition, triggers) in enumerate(tl_data):
    cell0 = tl_table.cell(i + 1, 0)
    format_table_cell(cell0, level, font_size=9, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell0, color_hex)
    format_table_cell(tl_table.cell(i + 1, 1), definition, font_size=9)
    format_table_cell(tl_table.cell(i + 1, 2), triggers, font_size=9)

tl_table.columns[0].width = Inches(1.0)
tl_table.columns[1].width = Inches(2.8)
tl_table.columns[2].width = Inches(3.0)
remove_table_borders(tl_table)

# Section: Important Reminders
guide_heading("Important Reminders")
guide_bullet("Flash alerts for breaking/critical events should be sent IMMEDIATELY via email and/or Slack — do not hold them for the morning SITREP.", bold_prefix="Flash Alerts: ")
guide_bullet("Save each completed report as: SITREP_YYYYMMDD.docx in the shared security drive.", bold_prefix="File Naming: ")
guide_bullet("Aim to distribute by 08:00 AM each business day.", bold_prefix="Timing: ")
guide_bullet("Review previous day's SITREP to update Open Items and track resolution of prior issues.", bold_prefix="Continuity: ")

# Section: Sources to Check Before Drafting
guide_heading("Sources to Check Before Drafting")
guide_bullet("Health Information Sharing and Analysis Center — sector-specific threat intelligence and alerts", bold_prefix="H-ISAC: ")
guide_bullet("Health Sector Cybersecurity Coordination Center — HHS threat briefs, analyst notes, and vulnerability bulletins", bold_prefix="HHS HC3: ")
guide_bullet("CISA.gov advisories, ICS-CERT alerts, and the Known Exploited Vulnerabilities (KEV) catalog", bold_prefix="CISA: ")
guide_bullet("FBI InfraGard portal — law enforcement threat intelligence for critical infrastructure", bold_prefix="FBI InfraGard: ")
guide_bullet("CrowdStrike, Palo Alto Unit 42, Microsoft Threat Intelligence, or other deployed vendor portals", bold_prefix="Vendor Intel: ")
guide_bullet("BleepingComputer, The Record, HealthITSecurity, SC Media, Dark Reading", bold_prefix="News Sources: ")
guide_bullet("Internal SIEM alerts, SOC shift reports, and helpdesk ticket trends from the previous 24 hours", bold_prefix="Internal Sources: ")

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
doc.save('/home/ubuntu/Daily_Cybersecurity_SITREP_Template.docx')
print("✅ Document saved successfully!")
